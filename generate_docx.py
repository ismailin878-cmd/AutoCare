import os
import sys
import subprocess

# Auto-install python-docx if not installed
try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx modulu bulunamadi. Yukleniyor...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Page setup - Margins (default is fine)
    
    # ----------------------------------------------------
    # TITLE PAGE
    # ----------------------------------------------------
    for _ in range(5):
        doc.add_paragraph() # spacing
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AUTOCARE ARAÇ BAKIM VE OTO SERVİS\nDİJİTAL ALTYAPI PROJESİ")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    for _ in range(6):
        doc.add_paragraph()
        
    p_student = doc.add_paragraph()
    p_student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = p_student.add_run("[Ad SOYAD]\n[Öğrenci No]\n\n")
    run_st.font.name = 'Arial'
    run_st.font.size = Pt(14)
    run_st.font.bold = True
    run_st.font.color.rgb = RGBColor(55, 65, 81)
    
    run_dept = p_student.add_run("BİLGİSAYAR TEKNOLOJİSİ VE BİLİŞİM SİSTEMLERİ BÖLÜMÜ\n")
    run_dept.font.name = 'Arial'
    run_dept.font.size = Pt(12)
    run_dept.font.bold = True
    
    run_course = p_student.add_run("BTS304 - Veritabanı Yönetim Sistemleri II\nFinal Ödevi\n")
    run_course.font.name = 'Arial'
    run_course.font.size = Pt(12)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_footer.add_run("2026\nBartın")
    run_foot.font.name = 'Arial'
    run_foot.font.size = Pt(11)
    run_foot.font.bold = True
    
    doc.add_page_break()
    
    # Helper to add standard headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(71, 85, 105)
        return h

    def add_body(text, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        
        # Add background or borders style if possible, or just Courier New
        r = p.add_run(text)
        r.font.name = 'Courier New'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(9, 10, 10)
        return p

    # ----------------------------------------------------
    # SECTION 1: SENARYO
    # ----------------------------------------------------
    add_heading_1("ADIM-1: Senaryo")
    
    add_body("Bu uygulama, \"AutoCare\" isimli modern bir araç bakım ve oto servis işletmesinin dijital altyapısını optimize etmek, operasyonel süreçlerini hızlandırmak ve veritabanı yönetimini kurumsal standartlara taşımak amacıyla hazırlanmıştır. İşletmenin günlük işleyişinde müşteri kayıtlarının düzgün tutulması, bu müşterilere ait araç bilgilerinin sisteme hatasız işlenmesi ve veri bütünlüğünün korunması hayati önem taşımaktadır.")
    
    add_body("Tasarladığımız veri tabanı ve otomasyon için belirlenen bazı kısıtlar ve kurallar aşağıda verilmiştir:", bold=True)
    
    add_bullet("Kayıtlı olmayan müşterilere araç tanımlanamaz, satış veya işlem gerçekleştirilemez.")
    add_bullet("Sisteme kaydedilmemiş araçlara herhangi bir bakım veya servis işlemi (yağ değişimi, balata değişimi vb.) uygulanamaz.")
    add_bullet("Servis işlemlerinde yedek parçalar için stok miktarı takibi yapılır. Stok miktarı 0 olan yedek parçalar işleme eklenmek istendiğinde veritabanı tetikleyicisi (Trigger) işlemi engeller ve hata fırlatır.")
    add_bullet("Bir servis işlemi gerçekleştirildiğinde kullanılan yedek parçaların stokları otomatik olarak 1 adet azaltılır. İşlemin iptal edilmesi veya silinmesi durumunda yedek parça stoğu otomatik 1 adet arttırılarak iade edilir.")
    add_bullet("Müşterinin güncel borç/alacak bakiyesi veritabanı fonksiyonu ile dinamik olarak hesaplanır. Müşteri Bakiyesi = (Toplam Ödemeleri) - (Müşterinin Araçlarına Yapılan Servis İşlemleri Toplamı) şeklinde hesaplanır.")
    add_bullet("Ödeme türleri sadece Nakit, Kredi Kartı ve Banka Havalesi olarak kabul edilir.")
    add_bullet("Hizmetlerin ve parçaların birimi Adet (stoklu) veya Saat (işçilik) olabilir.")

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 2: VARLIKLAR, İLİŞKİLER VE ER DIAGRAMI
    # ----------------------------------------------------
    add_heading_1("ADIM-2: Varlıklar, Nitelikler, İlişkiler ve ER Diagramı")
    
    add_heading_2("Sistemde Yer Alan Varlıklar ve Alanları")
    add_bullet("Müşteriler (musteri_id, musteri_ad, musteri_soyad, musteri_tel, musteri_mail, musteri_adres)")
    add_bullet("Araçlar (arac_id, musteri_id, arac_plaka, arac_marka, arac_model, arac_yil, arac_sasi_no)")
    add_bullet("Hizmetler/Ürünler (hizmet_id, hizmet_ad, hizmet_kategori, hizmet_fiyat, hizmet_stok, hizmet_birim, hizmet_detay)")
    add_bullet("Servis İşlemleri (islem_id, arac_id, hizmet_id, islem_tarih, islem_fiyat)")
    add_bullet("Ödemeler (odeme_id, musteri_id, odeme_tarih, odeme_tutar, odeme_tur, odeme_aciklama)")

    add_heading_2("Varlıklar Arası İlişkiler")
    add_bullet("Müşteri-Araç: Bir müşteriye ait birden fazla araç olabilir (1:N).")
    add_bullet("Araç-Servis İşlemi: Bir araç birden fazla kez servise girebilir ve işlem görebilir (1:N).")
    add_bullet("Hizmet-Servis İşlemi: Bir hizmet/yedek parça birden fazla servis işleminde kullanılabilir (1:N).")
    add_bullet("Müşteri-Ödeme: Bir müşteri birden fazla ödeme yapabilir (1:N).")

    add_heading_2("İlişkisel (Mantıksal) Şema")
    add_bullet("autocare_musteriler = {musteri_id, musteri_ad, musteri_soyad, musteri_tel, musteri_mail, musteri_adres}")
    add_bullet("autocare_araclar = {arac_id, +musteri_id, arac_plaka, arac_marka, arac_model, arac_yil, arac_sasi_no}")
    add_bullet("autocare_hizmetler = {hizmet_id, hizmet_ad, hizmet_kategori, hizmet_fiyat, hizmet_stok, hizmet_birim, hizmet_detay}")
    add_bullet("autocare_servis_islemleri = {islem_id, +arac_id, +hizmet_id, islem_tarih, islem_fiyat}")
    add_bullet("autocare_odemeler = {odeme_id, +musteri_id, odeme_tarih, odeme_tutar, odeme_tur, odeme_aciklama}")

    add_heading_2("ER Diyagramı")
    add_body("Aşağıdaki ER Diyagramı şeması, sistemin varlık ilişkilerini ve tabloların bağlantılarını görselleştirmektedir:")
    try:
        import os
        img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'er_diagram.png')
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=docx.shared.Inches(5.5))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.add_run("\nŞekil 1: AutoCare Veritabanı Varlık-İlişki (ER) Diyagramı").font.bold = True
        else:
            add_body("[ER Diyagramı görseli er_diagram.png bulunamadı, lütfen bu alana diyagramı ekleyiniz.]")
    except Exception as e:
        add_body(f"[ER Diyagramı eklenirken hata oluştu: {str(e)}]")

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 3: FİZİKSEL TASARIM
    # ----------------------------------------------------
    add_heading_1("ADIM-3: Fiziksel Yapı (SQL Kodları)")
    
    add_heading_2("Tablo Oluşturma SQL Kodları")
    
    sql_tables = """CREATE DATABASE autocare_db CHARACTER SET utf8mb4 COLLATE utf8mb4_0900_ai_ci;
USE autocare_db;

CREATE TABLE autocare_musteriler (
    musteri_id VARCHAR(64) NOT NULL,
    musteri_ad VARCHAR(64) NOT NULL,
    musteri_soyad VARCHAR(64) NOT NULL,
    musteri_tel VARCHAR(25) NOT NULL,
    musteri_mail VARCHAR(250) NOT NULL,
    musteri_adres VARCHAR(250) NOT NULL,
    PRIMARY KEY (musteri_id)
);

CREATE TABLE autocare_araclar (
    arac_id VARCHAR(64) NOT NULL,
    musteri_id VARCHAR(64) NOT NULL,
    arac_plaka VARCHAR(20) NOT NULL,
    arac_marka VARCHAR(50) NOT NULL,
    arac_model VARCHAR(50) NOT NULL,
    arac_yil INT NOT NULL,
    arac_sasi_no VARCHAR(50),
    PRIMARY KEY (arac_id),
    UNIQUE KEY (arac_plaka),
    UNIQUE KEY (arac_sasi_no),
    FOREIGN KEY (musteri_id) REFERENCES autocare_musteriler(musteri_id) ON DELETE CASCADE ON UPDATE CASCADE,
    CONSTRAINT chk_arac_yil CHECK (arac_yil >= 1900)
);

CREATE TABLE autocare_hizmetler (
    hizmet_id VARCHAR(64) NOT NULL,
    hizmet_ad VARCHAR(250) NOT NULL,
    hizmet_kategori VARCHAR(250) NOT NULL,
    hizmet_fiyat FLOAT NOT NULL,
    hizmet_stok FLOAT NOT NULL,
    hizmet_birim VARCHAR(16) NOT NULL,
    hizmet_detay VARCHAR(250) NOT NULL,
    PRIMARY KEY (hizmet_id),
    CONSTRAINT chk_hizmet_fiyat CHECK (hizmet_fiyat >= 0),
    CONSTRAINT chk_hizmet_stok CHECK (hizmet_stok >= 0)
);"""
    add_code(sql_tables)
    
    sql_tables_2 = """CREATE TABLE autocare_servis_islemleri (
    islem_id VARCHAR(64) NOT NULL,
    arac_id VARCHAR(64) NOT NULL,
    hizmet_id VARCHAR(64) NOT NULL,
    islem_tarih DATETIME NOT NULL,
    islem_fiyat FLOAT NOT NULL,
    PRIMARY KEY (islem_id),
    FOREIGN KEY (arac_id) REFERENCES autocare_araclar(arac_id) ON DELETE CASCADE ON UPDATE CASCADE,
    FOREIGN KEY (hizmet_id) REFERENCES autocare_hizmetler(hizmet_id) ON DELETE CASCADE ON UPDATE CASCADE,
    CONSTRAINT chk_islem_fiyat CHECK (islem_fiyat >= 0)
);

CREATE TABLE autocare_odemeler (
    odeme_id VARCHAR(64) NOT NULL,
    musteri_id VARCHAR(64) NOT NULL,
    odeme_tarih DATETIME NOT NULL,
    odeme_tutar FLOAT NOT NULL,
    odeme_tur VARCHAR(25) NOT NULL,
    odeme_aciklama VARCHAR(250) NOT NULL,
    PRIMARY KEY (odeme_id),
    FOREIGN KEY (musteri_id) REFERENCES autocare_musteriler(musteri_id) ON DELETE CASCADE ON UPDATE CASCADE,
    CONSTRAINT chk_odeme_tutar CHECK (odeme_tutar >= 0)
);"""
    add_code(sql_tables_2)

    doc.add_page_break()

    add_heading_2("Tetikleyiciler (Triggers)")
    sql_triggers = """DELIMITER //

-- 1. Stok Kontrolü Tetikleyicisi
CREATE TRIGGER tg_stok_kontrol
BEFORE INSERT ON autocare_servis_islemleri FOR EACH ROW
BEGIN
    DECLARE current_stok FLOAT;
    DECLARE item_birim VARCHAR(16);
    DECLARE hatamesaj VARCHAR(250);
    
    SELECT hizmet_stok, hizmet_birim INTO current_stok, item_birim
    FROM autocare_hizmetler WHERE hizmet_id = NEW.hizmet_id;
    
    IF (item_birim = 'Adet' AND current_stok < 1) THEN
        SET hatamesaj = CONCAT('HATA: Secilen yedek parca stokta kalmamistir!');
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = hatamesaj;
    END IF;
END; //

-- 2. Stok Azaltma Tetikleyicisi
CREATE TRIGGER tg_stok_azalt
AFTER INSERT ON autocare_servis_islemleri FOR EACH ROW
BEGIN
    DECLARE item_birim VARCHAR(16);
    SELECT hizmet_birim INTO item_birim FROM autocare_hizmetler WHERE hizmet_id = NEW.hizmet_id;
    
    IF (item_birim = 'Adet') THEN
        UPDATE autocare_hizmetler SET hizmet_stok = hizmet_stok - 1 WHERE hizmet_id = NEW.hizmet_id;
    END IF;
END; //

-- 3. Stok Geri Yükleme Tetikleyicisi
CREATE TRIGGER tg_stok_arttir
AFTER DELETE ON autocare_servis_islemleri FOR EACH ROW
BEGIN
    DECLARE item_birim VARCHAR(16);
    SELECT hizmet_birim INTO item_birim FROM autocare_hizmetler WHERE hizmet_id = OLD.hizmet_id;
    
    IF (item_birim = 'Adet') THEN
        UPDATE autocare_hizmetler SET hizmet_stok = hizmet_stok + 1 WHERE hizmet_id = OLD.hizmet_id;
    END IF;
END; //

DELIMITER ;"""
    add_code(sql_triggers)

    add_heading_2("Kullanıcı Tanımlı Fonksiyonlar (Functions)")
    sql_functions = """DELIMITER $$

CREATE FUNCTION fn_MusteriBakiye(p_musteri_id VARCHAR(64))
RETURNS FLOAT DETERMINISTIC
BEGIN
    DECLARE total_borc FLOAT DEFAULT 0;
    DECLARE total_odeme FLOAT DEFAULT 0;
    
    SELECT IFNULL(SUM(islem_fiyat), 0) INTO total_borc FROM autocare_servis_islemleri s
    JOIN autocare_araclar a ON s.arac_id = a.arac_id WHERE a.musteri_id = p_musteri_id;
    
    SELECT IFNULL(SUM(odeme_tutar), 0) INTO total_odeme FROM autocare_odemeler WHERE musteri_id = p_musteri_id;
    
    RETURN total_odeme - total_borc;
END $$

CREATE FUNCTION fn_AracServisSayisi(p_arac_id VARCHAR(64))
RETURNS INT DETERMINISTIC
BEGIN
    DECLARE service_count INT DEFAULT 0;
    SELECT COUNT(*) INTO service_count FROM autocare_servis_islemleri WHERE arac_id = p_arac_id;
    RETURN service_count;
END $$

DELIMITER ;"""
    add_code(sql_functions)

    doc.add_page_break()

    add_heading_2("Saklı Yordamlar (Stored Procedures - Örnekler)")
    add_body("Ödev kuralları gereği her tablo için SELECT, INSERT, UPDATE, DELETE işlemlerini yapan yordamlar yazılmıştır. Örnek olarak Müşteriler tablosuna ait yordamlar aşağıdadır:")
    
    sql_sp = """DELIMITER $$

CREATE PROCEDURE autocare_MusterilerHepsi()
BEGIN
    SELECT m.musteri_id AS ID, m.musteri_ad AS Adi, m.musteri_soyad AS Soyadi, m.musteri_tel AS Telefon, m.musteri_mail AS Mail, m.musteri_adres AS Adres, ( (SELECT IFNULL(SUM(o.odeme_tutar), 0) FROM autocare_odemeler o WHERE o.musteri_id = m.musteri_id) - (SELECT IFNULL(SUM(s.islem_fiyat), 0) FROM autocare_servis_islemleri s JOIN autocare_araclar a ON s.arac_id = a.arac_id WHERE a.musteri_id = m.musteri_id) ) AS Bakiye FROM autocare_musteriler m;
END $$

CREATE PROCEDURE autocare_MusteriEkle(p_id VARCHAR(64), p_ad VARCHAR(64), p_soy VARCHAR(64), p_tel VARCHAR(25), p_mail VARCHAR(250), p_adr VARCHAR(250))
BEGIN
    INSERT INTO autocare_musteriler(musteri_id, musteri_ad, musteri_soyad, musteri_tel, musteri_mail, musteri_adres) VALUES (p_id, p_ad, p_soy, p_tel, p_mail, p_adr);
END $$

CREATE PROCEDURE autocare_MusteriGuncelle(p_id VARCHAR(64), p_ad VARCHAR(64), p_soy VARCHAR(64), p_tel VARCHAR(25), p_mail VARCHAR(250), p_adr VARCHAR(250))
BEGIN
    UPDATE autocare_musteriler SET musteri_ad = p_ad, musteri_soyad = p_soy, musteri_tel = p_tel, musteri_mail = p_mail, musteri_adres = p_adr WHERE musteri_id = p_id;
END $$

CREATE PROCEDURE autocare_MusteriSil(p_id VARCHAR(64))
BEGIN
    DELETE FROM autocare_musteriler WHERE musteri_id = p_id;
END $$

CREATE PROCEDURE autocare_MusteriBul(p_filtre VARCHAR(64))
BEGIN
    SELECT m.musteri_id AS ID, m.musteri_ad AS Adi, m.musteri_soyad AS Soyadi, m.musteri_tel AS Telefon, m.musteri_mail AS Mail, m.musteri_adres AS Adres, ( (SELECT IFNULL(SUM(o.odeme_tutar), 0) FROM autocare_odemeler o WHERE o.musteri_id = m.musteri_id) - (SELECT IFNULL(SUM(s.islem_fiyat), 0) FROM autocare_servis_islemleri s JOIN autocare_araclar a ON s.arac_id = a.arac_id WHERE a.musteri_id = m.musteri_id) ) AS Bakiye FROM autocare_musteriler m WHERE m.musteri_id LIKE CONCAT('%', p_filtre, '%') OR m.musteri_ad LIKE CONCAT('%', p_filtre, '%') OR m.musteri_soyad LIKE CONCAT('%', p_filtre, '%') OR m.musteri_tel LIKE CONCAT('%', p_filtre, '%') OR m.musteri_mail LIKE CONCAT('%', p_filtre, '%') OR m.musteri_adres LIKE CONCAT('%', p_filtre, '%');
END $$

DELIMITER ;"""
    add_code(sql_sp)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 4: UYGULAMA GELİŞTİRME
    # ----------------------------------------------------
    add_heading_1("ADIM-4: Arayüz Tasarımı ve Uygulama Geliştirme")
    
    add_body("Uygulamamız, modern bir web otomasyon arayüzü olarak HTML5, CSS3 (Glassmorphism tasarımı) ve JavaScript (Vanilla JS) kullanılarak kodlanmıştır.")
    add_body("Arka planda (Backend) ise Node.js ve Express.js framework'ü kullanılmıştır. Uygulama mimarisi ders standartlarında N-Tier (N-Katmanlı) olarak kurulmuştur:")
    
    add_bullet("Presentation Layer (Sunum Katmanı): Tarayıcıda çalışan HTML/CSS/JS arayüzüdür. Kullanıcı form girdilerini doğrular ve sunucu API'sine istek gönderir.")
    add_bullet("Business Layer (İş Mantığı Katmanı): bl/ klasörü altındaki modüllerdir. Verilerin iş kurallarına uygunluğunu denetler, doğrulama yapar ve DAL katmanını çağırır.")
    add_bullet("Data Access Layer (Veri Erişim Katmanı): dal/ klasörü altındadır. KESİNLİKLE doğrudan SQL komutu barındırmaz. Sadece veritabanındaki Stored Procedure'leri tetikler.")

    add_heading_2("Veritabanı Erişim Güvenliği")
    add_body("Uygulamada SQL injection açıklarını önlemek ve ders kurallarına uymak amacıyla, sunucu tarafında hiçbir SQL sorgusu doğrudan çalıştırılmamaktadır. Tüm CRUD işlemleri veritabanı yordamları (Stored Procedures) aracılığıyla güvenli parametrelerle çağrılmaktadır.")

    add_heading_2("Ekran Görüntüleri ve Test Senaryoları")
    add_body("[BURAYA UYGULAMANIZIN EKRAN GÖRÜNTÜLERİNİ EKLEYİNİZ]")
    add_body("1. Dashboard: Genel ciro, tahsilat ve bakiye raporlarının gösterildiği arayüz.")
    add_body("2. Müşteri Kayıt: Yeni müşteri ekleme formu ve bakiyelerin dinamik gösterimi.")
    add_body("3. Tetikleyici Testi: Stokta 0 adet olan bir parçayı servis işlemine eklemek istediğinizde çıkan 'HATA: Secilen yedek parca stokta kalmamistir!' veritabanı trigger uyarı ekranı.")

    # Save to the user desktop or workspace
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    out_path = os.path.join(desktop_path, 'ad_soyad_rapor.docx')
    
    # Fallback to local workspace if Desktop doesn't exist or isn't writable
    try:
        doc.save(out_path)
        print(f"Rapor basariyla olusturuldu: {out_path}")
    except Exception as e:
        local_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ad_soyad_rapor.docx')
        doc.save(local_path)
        print(f"Masaustune kaydedilemedi. Calisma dizinine kaydedildi: {local_path}")

if __name__ == '__main__':
    create_report()
